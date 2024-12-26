import os
from docx import Document
from docx.enum.text import WD_BREAK

class DocumentArchiver:
    def __init__(self, output_dir, output_format, add_section_title=False, docx_in_docx_mode=False):
        """
        Initializes the DocumentArchiver.

        :param output_dir: Directory where the output files will be saved.
        :param output_format: Format of the output files ("txt" or "docx").
        :param add_section_title: Whether to add section titles in the output.
        :param docx_in_docx_mode: Whether to use advanced DOCX generation mode.
        """
        self.output_dir = output_dir
        self.output_format = output_format
        self.add_section_title = add_section_title
        self.docx_in_docx_mode = docx_in_docx_mode

    def archive_document(self, doc_path, sections, results, processor):
        """
        Archives the processed document in the specified format.

        :param doc_path: Path to the original document.
        :param sections: Parsed sections of the document.
        :param results: Processed results for each section.
        :param processor: The processor instance used for handling sections.
        """
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        full_name = f"{base_name}_{processor.output_suffix()}"

        if self.output_format == "txt":
            self._save_as_txt(full_name, results)
        elif self.output_format == "docx":
            self._save_as_docx(doc_path, sections, results, full_name)
        else:
            raise ValueError(f"Unrecognised output format: {self.output_format}")

    def _save_as_txt(self, full_name, results):
        if results:
            out_file = os.path.join(self.output_dir, f"{full_name}.txt")
            with open(out_file, 'w', encoding='utf-8') as f:
                for i, result in enumerate(results):
                    if self.add_section_title:
                        f.write(f"Section {i + 1}:\n")
                    f.write(result + "\n" + "=" * 40 + "\n")

    def _save_as_docx(self, doc_path, sections, results, full_name):
        out_file = os.path.join(self.output_dir, f"{full_name}.docx")
        if self.docx_in_docx_mode and self.output_format == "docx":
            self._generate_docx_advanced_mode(doc_path, sections, results, out_file)
        else:
            self._generate_docx(doc_path, sections, results, out_file)

    # Output docx for simple mode: just write section-by-section
    def _generate_docx(self, original_path, sections, results, output_path):
        doc = Document(original_path)
        text_map = self.map_sections(sections, results)
        for p in doc.paragraphs:
            original_text = p.text
            replaced = self.replace_text_in_sections(original_text, text_map)
            p.text = replaced
        doc.save(output_path)

    # Output docx for advanced mode: process a previously mapped set of dictionaries, allowing
    # paragraph-by-paragraph matching between original text and processed one    
    def _generate_docx_advanced_mode(self, original_path, sections, results, output_path):
        doc = Document(original_path)
        results_by_id = {r["id"]: r for r in results}  # Create a mapping of results by ID

        if len(sections) != len(results):
            raise ValueError("Sections and results lengths do not match.")

        def process_paragraph(paragraph, content, style_name):
            # Save original font settings
            original_font_name = None
            original_font_size = None
            if paragraph.runs:
                original_font_name = paragraph.runs[0].font.name
                original_font_size = paragraph.runs[0].font.size

            # Preserve non-text elements by clearing only text in existing runs
            if content.strip():
                if len(paragraph.runs) > 0:
                    paragraph.runs[0].text = content  # Modify the first run only
                    for run in paragraph.runs[1:]:
                        run.text = ''  # Clear all other runs
                else:
                    paragraph.add_run(content)  # Add content if no runs exist

            # Restore original font settings
            if paragraph.runs:
                run = paragraph.runs[0]
                if original_font_name:
                    run.font.name = original_font_name
                if original_font_size:
                    run.font.size = original_font_size

            # Apply styles if specified
            if style_name and style_name in [s.name for s in paragraph.part.styles]:
                paragraph.style = style_name

        # Track matched paragraphs by their IDs to prevent overwriting
        matched_ids = set()

        # Process main document paragraphs
        for idx, paragraph in enumerate(doc.paragraphs):
            section_id = f"main-{idx}"
            rendered_info = results_by_id.get(section_id)
            if rendered_info and section_id not in matched_ids:
                process_paragraph(paragraph, rendered_info.get("content", "").strip(), rendered_info.get("style_name", "").strip())
                matched_ids.add(section_id)

        # Process headers and footers
        for section in doc.sections:
            def process_section(paragraphs, section_type):
                for idx, paragraph in enumerate(paragraphs):
                    section_id = f"{section_type}-{idx}"
                    rendered_info = results_by_id.get(section_id)
                    if rendered_info and section_id not in matched_ids:
                        process_paragraph(paragraph, rendered_info.get("content", "").strip(), rendered_info.get("style_name", "").strip())
                        matched_ids.add(section_id)

            process_section(section.header.paragraphs, "header")
            process_section(section.footer.paragraphs, "footer")

        doc.save(output_path)

    def replace_text_in_sections(self, text, text_map):
        for original, translated in text_map.items():
            if original.strip() and original in text:
                text = text.replace(original, translated)
        return text

    def map_sections(self, sections, results):
        mapping = {}
        idx = 0
        for s in sections:
            mapping[s["content"]] = results[idx] if idx < len(results) else s["content"]
            idx += 1
        return mapping