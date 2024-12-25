#!/usr/bin/env python3

"""
Parses documents into structured sections based on headings and content.

- Supports `.docx`, plain text files, and `.pdf`.
- Merges sections with low word count below a configurable threshold (for DOCX).
- Customizable heading styles determine section boundaries (for DOCX).
- PDF documents are split by page, each becoming its own section.
- Can return sections with or without title, depending on the param in main
"""

import os
from docx import Document
from PyPDF2 import PdfReader
import tiktoken

class DocumentParser:
    def __init__(self, heading_styles, min_word_threshold=2):
        """
        :param heading_styles: A set or list of style names considered headings in DOCX.
        :param min_word_threshold: Sections below this word count will be merged with the next.
        """
        self.heading_styles = heading_styles
        self.min_word_threshold = min_word_threshold

    def parse_document(self, file_path, docx_in_docx_mode=False):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            if docx_in_docx_mode:
                return self._parse_docx_by_paragraph(file_path)
            else:
                return self._parse_docx(file_path)
        elif ext == ".pdf":
            return self._parse_pdf(file_path)
        else:
            return self._parse_text(file_path)

    def _parse_text(self, file_path):
        """
        Parses a plain text file as a single section.
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read().strip()
        if not text:
            return []
        sections = [{"title": "Document", "content": text}]
        return self._split_sections_if_needed(sections)
    
    def _parse_docx(self, file_path):
        """
        Parses a DOCX file by headings, ensuring titles merge with the following paragraphs
        and including intro if no headings are present initially.
        """
        document = Document(file_path)
        paragraphs = document.paragraphs

        sections = []
        current_section = []
        intro_section = []
        is_intro = True

        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue  # Skip empty paragraphs

            style_name = paragraph.style.name if paragraph.style else ""
            if style_name in self.heading_styles:  # New section start
                if current_section:
                    sections.append(current_section)  
                current_section = [text]  
                is_intro = False
            else:
                if is_intro:
                    intro_section.append(text) 
                elif current_section:
                    current_section.append(text)  

        if current_section:
            sections.append(current_section)
        if intro_section:  
            sections.insert(0, ["Introduction", *intro_section])  # Prepend intro section if it exists

        merged_sections = []
        i = 0
        while i < len(sections):
            section = sections[i]
            title = section[0] if section else "Untitled"
            content = "\n".join(section[1:]).strip()
            word_count = len(content.split())
            j = i

            while word_count < self.min_word_threshold and j + 1 < len(sections):
                j += 1
                next_sec = sections[j]
                content += "\n" + "\n".join(next_sec)
                word_count = len(content.split())

            merged_sections.append({"title": title, "content": content})
            i = j + 1

        return self._split_sections_if_needed(merged_sections)

    def _parse_docx_by_paragraph(self, file_path):
        document = Document(file_path)
        sections = []
        for idx, paragraph in enumerate(document.paragraphs):
            sections.append({
                "id": idx,  # Add unique identifier for each section
                "title": f"Paragraph {idx+1}",
                "content": paragraph.text,
                "style_name": paragraph.style.name if paragraph.style else None
            })
        return sections

    def _parse_pdf(self, file_path):
        """
        Parses a PDF file by splitting each page into its own section.
        """
        reader = PdfReader(file_path)
        sections = []

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            sections.append({"title": f"PDF Page {i + 1}", "content": page_text or ""})
        return self._split_sections_if_needed(sections)

    def _split_sections_if_needed(self, sections):
        """
        Splits sections if their content exceeds the maximum token limit.
        """
        max_tokens = self._calculate_max_tokens()
        split_sections = []

        for section in sections:
            content = section["content"].strip()
            if self._calculate_tokens(content) > max_tokens:
                split_sections.extend(self._smart_split(section))
            else:
                split_sections.append(section)

        return split_sections

    def _calculate_max_tokens(self):
        """
        Calculates the maximum reasonable tokens for one section.
        """
        return 7000  # TODO: dynamic logic depending on model type?

    def _calculate_tokens(self, text):
        """
        Calculates the number of tokens in a given text.
        """
        encoding = tiktoken.encoding_for_model("gpt-4")
        return len(encoding.encode(text))

    def _smart_split(self, section):
        """
        Splits a section into smaller sections without breaking points or sentences.
        """
        max_tokens = self._calculate_max_tokens()
        content = section["content"].strip()
        title = section["title"]

        sentences = content.split('.')
        split_sections = []
        current_chunk = ""

        for sentence in sentences:
            if self._calculate_tokens(current_chunk + sentence + '.') > max_tokens:
                split_sections.append({"title": title, "content": current_chunk.strip()})
                current_chunk = sentence + '.'
            else:
                current_chunk += sentence + '.'

        if current_chunk:
            split_sections.append({"title": title, "content": current_chunk.strip()})

        return split_sections
