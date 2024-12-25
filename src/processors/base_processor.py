#!/usr/bin/env python3

"""
BaseProcessor class provides a generic framework for processing document sections.

- Includes utilities for generating updated `.docx` files by replacing text in original documents.
- Intended to be extended by specific processors like reporters or reviewers.
"""

import os
import copy
from docx import Document
from docx.enum.text import WD_BREAK

class BaseProcessor:
    def __init__(self, client, processor_parameters):
        self.client = client
        self.processor_parameters = processor_parameters

    # Override this to declare sections processing logic
    def process_sections(self, sections):
        return sections

    # Override this to declare section postprocessing logic
    def postprocess(self, response, section):
        return response

    def output_suffix(self):
        return "processed"

    # File utility to output a docx
    def generate_docx(self, original_path, sections, results, output_path):
        doc = Document(original_path)
        text_map = self.map_sections(sections, results)
        for p in doc.paragraphs:
            original_text = p.text
            replaced = self.replace_text_in_sections(original_text, text_map)
            p.text = replaced
        doc.save(output_path)

    def generate_docx_advanced_mode(self, original_path, sections, results, output_path):
        doc = Document(original_path)
        results_by_id = {r["id"]: r for r in results}  # Create a mapping of results by ID

        if len(sections) != len(results):
            raise ValueError("Sections and results lengths do not match.")

        def process_paragraph(paragraph, content, style_name):
            # Save original font settings
            original_font_name = None
            original_font_size = None
            if paragraph.runs:
                original_font_name = paragraph.runs[0].font.name
                original_font_size = paragraph.runs[0].font.size

            # Preserve non-text elements by clearing only text in existing runs
            if content.strip():
                if len(paragraph.runs) > 0:
                    paragraph.runs[0].text = content  # Modify the first run only
                    for run in paragraph.runs[1:]:
                        run.text = ''  # Clear all other runs
                else:
                    paragraph.add_run(content)  # Add content if no runs exist

            # Restore original font settings
            if paragraph.runs:
                run = paragraph.runs[0]
                if original_font_name:
                    run.font.name = original_font_name
                if original_font_size:
                    run.font.size = original_font_size

            # Apply styles if specified
            if style_name and style_name in [s.name for s in paragraph.part.styles]:
                paragraph.style = style_name

        # Track matched paragraphs by their IDs to prevent overwriting
        matched_ids = set()

        # Process main document paragraphs
        for idx, paragraph in enumerate(doc.paragraphs):
            section_id = f"main-{idx}"
            rendered_info = results_by_id.get(section_id)
            if rendered_info and section_id not in matched_ids:
                process_paragraph(paragraph, rendered_info.get("content", "").strip(), rendered_info.get("style_name", "").strip())
                matched_ids.add(section_id)

        # Process headers and footers
        for section in doc.sections:
            def process_section(paragraphs, section_type):
                for idx, paragraph in enumerate(paragraphs):
                    section_id = f"{section_type}-{idx}"
                    rendered_info = results_by_id.get(section_id)
                    if rendered_info and section_id not in matched_ids:
                        process_paragraph(paragraph, rendered_info.get("content", "").strip(), rendered_info.get("style_name", "").strip())
                        matched_ids.add(section_id)

            process_section(section.header.paragraphs, "header")
            process_section(section.footer.paragraphs, "footer")

        doc.save(output_path)

    def map_sections(self, sections, results):
        return {section["content"]: result for section, result in zip(sections, results)}

    def replace_text_in_sections(self, text, text_map):
        for original, translated in text_map.items():
            if original.strip() and original in text:
                text = text.replace(original, translated)
        return text

    def map_sections(self, sections, results):
        mapping = {}
        idx = 0
        for s in sections:
            mapping[s["content"]] = results[idx] if idx < len(results) else s["content"]
            idx += 1
        return mapping

    def replace_text_in_sections(self, text, text_map):
        for k, v in text_map.items():
            if k.strip() and k in text:
                text = text.replace(k, v)
        return text